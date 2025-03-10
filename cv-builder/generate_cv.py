import nltk
from nltk.corpus import wordnet
from docx import Document

# Download necessary NLTK data
nltk.download('wordnet')

def generate_synonyms(word):
    synonyms = []
    for syn in wordnet.synsets(word):
        for lemma in syn.lemmas():
            synonyms.append(lemma.name())
    return set(synonyms)

def generate_content(prompt):
    keywords = prompt.split()
    synonyms = {word: generate_synonyms(word) for word in keywords}
    return ' '.join([', '.join(synonyms[word]) for word in keywords])

def create_cv(template_path, output_path, personal_details, experience, education, skills, custom_sections):
    doc = Document(template_path)

    # Replace placeholders with personal details
    for paragraph in doc.paragraphs:
        for key, value in personal_details.items():
            if f'{{{{ {key} }}}}' in paragraph.text:
                paragraph.text = paragraph.text.replace(f'{{{{ {key} }}}}', value)

    # Add experience
    experience_section = doc.add_heading('Experience', level=1)
    for job in experience:
        doc.add_heading(job['title'], level=2)
        doc.add_paragraph(job['description'])

    # Add education
    education_section = doc.add_heading('Education', level=1)
    for edu in education:
        doc.add_heading(edu['degree'], level=2)
        doc.add_paragraph(edu['institution'])

    # Add skills
    skills_section = doc.add_heading('Skills', level=1)
    doc.add_paragraph(', '.join(skills))

    # Add custom sections
    for section in custom_sections:
        doc.add_heading(section['title'], level=1)
        doc.add_paragraph(section['content'])

    doc.save(output_path)

# Example usage
personal_details = {
    "name": "John Doe",
    "contact_info": "john.doe@example.com"
}

experience = [
    {"title": "Software Engineer", "description": generate_content("Developed web applications and analyzed data trends.")}
]

education = [
    {"degree": "B.Sc. Computer Science", "institution": "University X"},
    {"degree": "M.Sc. Data Science", "institution": "University Y"}
]

skills = ["Python", "Docker", "Machine Learning"]

custom_sections = [
    {"title": "Projects", "content": "Project A, Project B"},
    {"title": "Certifications", "content": "Certification A, Certification B"}
]

create_cv('cv_template.docx', 'personalized_cv.docx', personal_details, experience, education, skills, custom_sections)
